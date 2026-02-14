"""
Word Document (DOCX) Loader

PURPOSE:
Loads and extracts text from Microsoft Word documents (.docx, .doc).

WHY DOCX SUPPORT:
- Many business documents are in Word format
- Contracts, proposals, policies often in DOCX
- Need to extract text while preserving structure

LIBRARY USED: python-docx
- Industry standard for DOCX files
- Extracts paragraphs, tables, headers, footers
- Preserves document structure
"""

from typing import List
from pathlib import Path

from .base_loader import BaseDocumentLoader, Document

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class DOCXLoader(BaseDocumentLoader):
    """
    Loads Microsoft Word documents (.docx).
    
    WHAT IT DOES:
    - Extracts text from paragraphs
    - Extracts text from tables
    - Preserves document structure
    - Handles headers and footers (optional)
    
    USAGE:
        loader = DOCXLoader("contract.docx")
        documents = loader.load()
        
        for doc in documents:
            print(doc.content)
    """
    
    def __init__(self, file_path: str, include_tables: bool = True):
        """
        Initialize DOCX loader.
        
        Args:
            file_path: Path to DOCX file
            include_tables: Whether to extract text from tables
        """
        super().__init__(file_path)
        
        # Validate it's a DOCX file
        if self.file_path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(
                f"Expected DOCX/DOC file, got: {self.file_path.suffix}\n"
                f"Use DOCXLoader only for .docx or .doc files"
            )
        
        self.include_tables = include_tables
    
    
    def load(self) -> List[Document]:
        """
        Load DOCX and extract all text.
        
        STRATEGY:
        - Extract all paragraphs
        - Extract tables if enabled
        - Combine into single document (or split by sections)
        
        RETURNS:
            List with one Document object containing all text
        
        NOTE: Unlike PDFs, DOCX files don't have "pages"
              We return one document with all content
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx not installed. Install with:\n"
                "  pip install python-docx"
            )
        
        print(f"📖 Loading DOCX: {self.file_path.name}")
        
        # Load the document
        doc = DocxDocument(self.file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Only include non-empty paragraphs
                paragraphs.append(text)
        
        print(f"   Extracted {len(paragraphs)} paragraphs")
        
        # Extract text from tables (if enabled)
        table_text = []
        if self.include_tables and doc.tables:
            for table_num, table in enumerate(doc.tables, start=1):
                table_content = self._extract_table_text(table)
                if table_content:
                    table_text.append(f"\n[Table {table_num}]\n{table_content}")
            
            print(f"   Extracted {len(doc.tables)} tables")
        
        # Combine all text
        all_text = "\n\n".join(paragraphs)
        if table_text:
            all_text += "\n\n" + "\n".join(table_text)
        
        # Clean the text
        cleaned_text = self.clean_text(all_text)
        
        # Create Document object
        document = Document(
            content=cleaned_text,
            metadata={
                "source": self.file_path.name,
                "file_path": str(self.file_path),
                "num_paragraphs": len(paragraphs),
                "num_tables": len(doc.tables) if self.include_tables else 0,
                "loader": "docx"
            }
        )
        
        print(f"✓ Extracted {len(cleaned_text)} characters")
        
        return [document]  # Return as list for consistency with other loaders
    
    
    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table.
        
        WHY SEPARATE METHOD:
        - Tables have complex structure (rows, cells)
        - Need to preserve table layout
        - Reusable for multiple tables
        
        Args:
            table: python-docx Table object
        
        RETURNS:
            Formatted table text
        """
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Join cells with | separator (markdown-style)
            row_text = " | ".join(cells)
            if row_text.strip():
                rows.append(row_text)
        
        return "\n".join(rows)


# ============================================================================
# MAIN EXECUTION (for testing)
# ============================================================================

if __name__ == "__main__":
    """
    Test the DOCX loader.
    
    Run: python src/loaders/docx_loader.py
    """
    import sys
    from src.config import RAW_DATA_DIR
    
    print("\n" + "="*60)
    print("TESTING DOCX LOADER")
    print("="*60 + "\n")
    
    # Look for DOCX files in data/raw/
    docx_files = list(RAW_DATA_DIR.glob("*.docx")) + list(RAW_DATA_DIR.glob("*.doc"))
    
    if not docx_files:
        print(f"❌ No DOCX files found in {RAW_DATA_DIR}")
        print(f"\nPlease place a .docx file there and try again.")
        print(f"\nAlternatively, create a test file:")
        print(f"  1. Open Microsoft Word")
        print(f"  2. Type some text")
        print(f"  3. Save as test.docx in {RAW_DATA_DIR}")
        sys.exit(1)
    
    # Test with first DOCX found
    test_file = docx_files[0]
    print(f"Testing with: {test_file.name}\n")
    
    try:
        # Load DOCX
        loader = DOCXLoader(str(test_file))
        documents = loader.load()
        
        # Show statistics
        print(f"\n📊 STATISTICS:")
        print(f"  Documents: {len(documents)}")
        
        if documents:
            doc = documents[0]
            print(f"  Total characters: {len(doc.content):,}")
            print(f"  Paragraphs: {doc.metadata['num_paragraphs']}")
            print(f"  Tables: {doc.metadata['num_tables']}")
            
            # Show sample
            print(f"\n📄 SAMPLE CONTENT:")
            print("-" * 60)
            print(doc.content[:500] + "...")
            print("-" * 60)
            
            print(f"\n📋 METADATA:")
            for key, value in doc.metadata.items():
                print(f"  {key}: {value}")
        
        print("\n✓ DOCX loader test successful!")
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
